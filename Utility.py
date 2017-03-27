__author__ = 'Alexander'

'''
functions to load and save words and count
'''
import pickle
def save_dict(obj, name):
    with open('resources/'+name+'.pkl','wb') as file:
        pickle.dump(obj, file, pickle.HIGHEST_PROTOCOL)
    print("saved")

def load_dict(name):
    with open('resources/'+name+'.pkl','rb') as file:
        return pickle.load(file)
    print("loaded")

'''
function to clear words from word that have no sense
'''
import re
def to_words(content, words):
    letters_only = re.sub("[^a-zA-Z]", " ", content)
    text = letters_only.lower().split()
    # stops = set(stopwords.words("english"))
    meaningful_words = [w for w in text if w not in words]
    return( " ".join( meaningful_words ))

from nltk.corpus import stopwords
def useful_words(content, words):
    letters_only = re.sub("[^a-zA-Z]", " ", content)
    text = letters_only.lower().split()
    # text = content.split()
    meaningful_words = [w for w in text if w in words]
    return( " ".join( meaningful_words ))

'''
collect all words and save it
'''
import pandas as pd
from nltk.corpus import stopwords
from time import time
from sklearn.feature_extraction.text import CountVectorizer
def collect_words_and_save(name):
    stops = set(stopwords.words("english"))
    df = pd.read_csv('./resources/Combined_News_DJIA.csv')
    differ_words = []
    t0 = time()
    tokenizer = CountVectorizer().build_tokenizer()
    for row in range(0, len(df.index)):
        title = ' '.join(str(x) for x in df.iloc[row, 2:27])
        print("done in %0.3fs." % (time() - t0))
        differ_words += tokenizer(title.lower())
    print("done for")
    words = set(differ_words)
    print("done set")

    # need 20 min to calculate for all data better to use pickle functions...
    dic = []
    t0 = time()
    for x in words:
        if x not in stops:
            dic.append([x, differ_words.count(x)])
            print("done in %0.3fs." % (time() - t0))
    print("done dict")

    save_dict(dic, name)
    return dic

'''
function to load data frame from file
'''
def manage_df_and_save(output, df, stops):
    Cleared = []
    for i in range(len(df)):
        Cleared.append(to_words(df.loc[i,'Combined'], stops))
    result = pd.DataFrame({'Label': df['Label'].tolist(),
                                'Cleared': Cleared})
    print("done managing")
    save_dict(result, output)
    return result

'''
function for import all data in file
'''
def import_in_file(words, Model, Accuracy, Crosstable, Plots, proportions):
    from docx import Document
    print("Saving to file...")
    fileName = "Project Information.docx"
    doc = open(fileName, 'rb')
    document = Document(doc)
    document.add_heading("Statistics for {} features".format(len(words)))
    #import statistics
    # document.add_heading("Word count statistics")
    # table = document.add_table(rows=1, cols=2)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Word'
    # hdr_cells[1].text = 'Count'
    # size = 10 #count rows to print
    # for i in range(0, size):
    #     row_cells = table.add_row().cells
    #     word = words[i]
    #     count = useful.loc[useful['Word'] == word].iat[0,1]
    #     row_cells[0].text = str(word)
    #     row_cells[1].text = str(count)
    # document.add_picture('count.png')
    # document.add_picture('wordcloud1.png')
    # document.add_picture('wordcloud2.png')

    #import models result
    for i in range(len(Model)):
        model = Model[i]
        accuracy = Accuracy[i]
        ct = Crosstable[i]
        plot = Plots[i]

        print("{} model".format(model))
        document.add_heading(model)

        if ct.shape[0]==2 and ct.shape[1]==2:
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Actual'
            hdr_cells[1].text = 'Predicted down'
            hdr_cells[2].text = 'Predicted up'
            row_cells = table.add_row().cells
            row_cells[0].text = "down"
            row_cells[1].text = str(round(ct[0][0]*proportions, 3))
            row_cells[2].text = str(round(ct[1][0]*proportions, 3))
            row_cells = table.add_row().cells
            row_cells[0].text = "up"
            row_cells[1].text = str(round(ct[0][1]*(1-proportions), 3))
            row_cells[2].text = str(round(ct[1][1]*(1-proportions), 3))

        # print("saving plot")
        # document.add_picture(plot)

        document.add_paragraph('Accuracy of '+str(model)+' is '+str(round(accuracy, 3)))

        # if (model == "LogisticRegression"):
        #     coef = Coefficients[i]
        #     table = document.add_table(rows=1, cols=2)
        #     hdr_cells = table.rows[0].cells
        #     hdr_cells[0].text = 'Coefficient'
        #     hdr_cells[1].text = 'Word'
        #     size = 5
        #     for j in range(0, size):
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = str(round(coef['Coefficient'][j], 3))
        #         row_cells[1].text = str(round(coef['Word'][j], 3))
        #     for j in range(size, 1):
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = str(round(coef['Coefficient'][-j], 3))
        #         row_cells[1].text = str(round(coef['Word'][-j], 3))
        # document.add_picture(plot)
    document.save(fileName)
    doc.close()
    print("done")